import docx
from docx.shared import Pt

src = 'MasterHelp.docx'
d = docx.Document(src)

print('== Sections / Page set-up ==')
for i, sec in enumerate(d.sections):
    print(f'  section {i}: page {sec.page_width.emu}x{sec.page_height.emu} EMU '
          f'( {sec.page_width.mm:.1f}mm x {sec.page_height.mm:.1f}mm )')

print('\n== Paragraph count ==', len(d.paragraphs))
print('== Table count ==', len(d.tables))
print('== Inline shapes (images) ==', len(d.inline_shapes))

print('\n== Style names used (unique) ==')
styles_used = {}
for p in d.paragraphs:
    s = p.style.name if p.style else 'None'
    styles_used[s] = styles_used.get(s, 0) + 1
for k, v in sorted(styles_used.items(), key=lambda kv: -kv[1]):
    print(f'  {k}: {v}')

print('\n== Headings (first 80) ==')
hdr_count = 0
for p in d.paragraphs:
    s = p.style.name if p.style else ''
    if s.startswith('Heading') or s.startswith('Title'):
        text = (p.text or '').strip()
        if text:
            print(f'  [{s}] {text[:120]}')
            hdr_count += 1
            if hdr_count >= 80:
                print('  ... (truncated)')
                break

print('\n== First 5 paragraphs (raw) ==')
for i, p in enumerate(d.paragraphs[:5]):
    print(f'  {i:3d} [{p.style.name}]: {p.text[:160]!r}')

print('\n== Total characters (all paragraphs) ==')
total_chars = sum(len(p.text) for p in d.paragraphs)
print(' ', total_chars, 'characters across', len(d.paragraphs), 'paragraphs')
