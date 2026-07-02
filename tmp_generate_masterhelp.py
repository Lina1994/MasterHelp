"""Generate the updated MasterHelp.docx end-to-end.

Produces a comprehensive reference manual of every feature/tool/option of
the MasterHelp application. Uses the same style palette as the previous
document (Title, Heading 1/2/3, List Bullet, tables) and expands each
chapter to cover features added in recent sprints (cartas, scenes,
recurring actor, audio orchestrator, etc.).
"""

import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = 'MasterHelp.docx'


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_table(doc, headers, rows, header_fill='1F4E79', bold_first=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(hdr_cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            if bold_first and i == 0:
                for p in cells[i].paragraphs:
                    for run in p.runs:
                        run.bold = True
    return table


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def p(doc, text, italic=False, bold=False, size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return para


def bullet(doc, text, level=0):
    return doc.add_paragraph(text, style='List Bullet')


def code_block(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    return para


# ---------------------------------------------------------------------------
# Document scaffold
# ---------------------------------------------------------------------------

doc = Document()

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Title page ---------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MasterHelp')
run.bold = True
run.font.size = Pt(36)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Guía completa de usuario y referencia de funcionalidades')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x36, 0x6F, 0xB1)

v = doc.add_paragraph()
v.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = v.add_run('Edición generada automáticamente · ' + datetime.date.today().isoformat())
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph().add_run().add_break()

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = intro.add_run(
    'MasterHelp es una aplicación de escritorio + web para dirigir campañas de '
    'D&D 5e y manuales compatibles. Este documento describe TODAS las opciones, '
    'herramientas y funcionalidades disponibles, agrupadas por área de uso. '
    'El texto está organizado en 11 capítulos y 4 apéndices.'
)
r.italic = True
r.font.size = Pt(11)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------

h1(doc, 'Índice')
toc = [
    ('Capítulo 1 · Introducción y conceptos básicos', '1'),
    ('Capítulo 2 · Primeros pasos', '8'),
    ('Capítulo 3 · HOME (Ajustes)', '13'),
    ('Capítulo 4 · Pantallas secundarias (Mapas, Skyline, Proyección)', '16'),
    ('Capítulo 5 · Sidebar y tira del Master', '22'),
    ('Capítulo 6 · Herramientas por módulo', '26'),
    ('Capítulo 7 · Atajos (Shortcuts)', '53'),
    ('Capítulo 8 · Iconos y ventanas flotantes Skyline', '55'),
    ('Capítulo 9 · Funcionalidades adicionales (Importación, Exportación, etc.)', '57'),
    ('Capítulo 10 · Acceso web, Electron y multi-dispositivo', '60'),
    ('Capítulo 11 · Roles de usuario: Master vs Jugador', '63'),
    ('Apéndice A · Mapa de rutas (frontend router)', '65'),
    ('Apéndice B · Superficie de la API (endpoints backend)', '70'),
    ('Apéndice C · Stack tecnológico', '78'),
    ('Apéndice D · Glosario y atajos globales del teclado', '83'),
]
for label, page in toc:
    para = doc.add_paragraph(style='List Bullet')
    r = para.add_run(f'{label}  ……  pág. {page}')
    r.font.size = Pt(11)
doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 1
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 1 · Introducción y conceptos básicos')
h2(doc, '1.1 ¿Qué es MasterHelp?')
p(doc,
  'MasterHelp es el compañero digital del Dungeon Master. Reúne en una sola app '
  'la gestión de campañas, manuales, mazmorras, criaturas, hechizos, personajes, '
  'misiones, tiendas, diario de sesiones, worldpedia (wiki de campaña), '
  'cartas imprimibles, soundtrack, efectos de sonido, escenas con síntesis de voz '
  'y proyección para los jugadores. Se ejecuta como aplicación Electron en escritorio '
  'y como SPA en navegador web; ambas comparten la misma base de código.')

h2(doc, '1.2 Roles de Usuario')
bullet(doc, 'Master (DM): control total sobre la campaña, contenido y proyección.')
bullet(doc, 'Player (Jugador): acceso de solo lectura a las ventanas de proyección y a su propio personaje.')
bullet(doc, 'Invitado: pendiente de aceptar invitación para unirse a una campaña.')

h2(doc, '1.3 Conceptos Fundamentales')
bullet(doc, 'Campaña: contenedor raíz. Aísla contenido, permisos y configuración por mesa de juego.')
bullet(doc, 'Manuales: bases de datos de reglas (D&D 5e por defecto y manuales personalizados) que proveen hechizos, clases, monstruos, etc.')
bullet(doc, 'Mapas: superficies rectangulares con imagen y skyline opcional por momento del día (amanecer, día, atardecer, noche).')
bullet(doc, 'Encuentro: combate temporal con iniciativa, participantes y notas. Se proyecta en ventana secundaria para los jugadores.')
bullet(doc, 'Escena: composición multimedia reproducible (audio, video, narración TTS) que el Master dispara manualmente.')
bullet(doc, 'Carta: objeto imprimible generado a partir de una plantilla + un personaje/hechizo/rasgo/dote.')
bullet(doc, 'Worldpedia: wiki jerárquica de la campaña con notas Markdown y backlinks.')
bullet(doc, 'Diario: registro cronológico de la campaña con calendario personalizado, sesiones y entradas diarias.')
bullet(doc, 'Soundtrack: listas de reproducción y canciones por campaña con historial de uso.')

h2(doc, '1.4 Stack Resumido')
p(doc, 'MasterHelp bebe del siguiente stack (ver Apéndice C para detalle).')
bullet(doc, 'Frontend: React 18 + Vite + Material UI 7 + i18next (es/en).')
bullet(doc, 'Backend: NestJS 10 + TypeORM sobre SQLite + Passport JWT.')
bullet(doc, 'Sincerización entre ventanas: BroadcastChannel + localStorage + REST polling.')
bullet(doc, 'Audio: Web Audio API nativa + reproducción cruzada.')
bullet(doc, 'Empaquetado: Electron 38 con actualizaciones automáticas (electron-updater).')
bullet(doc, 'Tipos compartidos en TypeScript.')

h2(doc, '1.5 Arquitectura general')
p(doc, 'Las dos ventanas de proyección (mapa y skyline) son páginas del propio router '
  '(/projection/maps y /projection/skyline), abiertas como BrowserWindow adicionales en '
  'Electron o como pestañas del navegador en web. Los endpoints /campaigns/projection/* son '
  'públicos (sin JWT) y devuelven solo el estado que los jugadores pueden ver. La ventana del '
  'master publica cambios sobre el estado de batalla, tiempo del día, niebla, grid y '
  'soundtrack; las ventanas de proyección se sincronizan vía polling y BroadcastChannel.')
doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 2
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 2 · Primeros pasos')
h2(doc, '2.1 Registro e Inicio de Sesión')
h3(doc, 'Registro')
bullet(doc, 'Web: /register.')
bullet(doc, 'Campos: email, contraseña (mín. 8 caracteres), nombre de usuario opcional, idioma preferido.')
bullet(doc, 'Verificación: el backend envía un correo de bienvenida vía Resend.')
bullet(doc, 'Tras registrar, el usuario queda autenticado con JWT en localStorage["access_token"].')
h3(doc, 'Login')
bullet(doc, 'Web: /login.')
bullet(doc, 'Persistencia: el token JWT caduca a las 24 h; el cliente intenta renovarlo silenciosamente antes de expirar.')
h3(doc, 'Recuperación de contraseña')
bullet(doc, 'Web: /forgot-password. Se envía un email con un enlace firmado.')
bullet(doc, 'Web: /reset-password recibe el token y permite definir contraseña nueva.')
h3(doc, 'Cambio de contraseña')
bullet(doc, 'Ruta autenticada: /change-password.')
bullet(doc, 'Requiere contraseña actual y nueva (mín. 8 caracteres).')
h3(doc, 'Eliminar cuenta')
bullet(doc, 'Ruta: /delete-account. Acción irreversible. Borra personajes, campañas propias y manuales personalizados.')

h2(doc, '2.2 Idioma y Tema')
bullet(doc, 'Selector de idioma (es | en) en HOME → Ajustes. Cambia toda la UI y los formatos de fecha.')
bullet(doc, 'Tema claro/oscuro disponible vía ThemeContext; el switch está en HOME → Tema.')
bullet(doc, 'Combinación preferida (idioma + tema) se guarda en /users/me/preferences.')

h2(doc, '2.3 Selección y Activación de Campaña')
bullet(doc, 'Tras autenticarse, el sidebar muestra todas las campañas visibles.')
bullet(doc, 'Al pulsar sobre una campaña, su id persiste en localStorage["activeCampaignId"] y el contexto global ActiveCampaignContext lo expone al resto de la app.')
bullet(doc, 'El banner del sidebar cambia al nombre e imagen de la campaña activa.')
bullet(doc, 'Si la campaña activa es eliminada, app fallback a /campaigns.')

h2(doc, '2.4 Invitar jugadores')
bullet(doc, 'Desde /campaigns → "Ajustes de campaña" → "Invitar".')
bullet(doc, 'Formulario: email (≤ 255 caracteres) y nombre de usuario opcional (≤ 50 caracteres).')
bullet(doc, 'Backend crea un registro de CampaignPlayer con status=pending.')
bullet(doc, 'Email enviado con magic link hacia /login?invitation=ID.')
bullet(doc, 'Al iniciar sesión, el destinatario ve la invitación en el banner superior (InvitationsList) y puede aceptar, declinar, o re-responder.')
doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 3
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 3 · HOME (Ajustes)')
p(doc, 'La página / (HomePage) monta un panel de SettingsSection con varios sub-paneles.')

h2(doc, '3.1 Actualizador automático')
bullet(doc, 'UpdateChecker se monta en cada Home y consulta la versión publicada cada 24 h.')
bullet(doc, 'Si hay actualización mayor, aparece un cuadro de diálogo que permite descargar e instalar antes del próximo arranque.')
bullet(doc, 'En entornos web, UpdateChecker muestra sólo mensaje informativo.')

h2(doc, '3.2 Personalización del sidebar')
bullet(doc, 'SidebarSettings permite elegir qué módulos (de los 15 predefinidos) son visibles.')
bullet(doc, 'Cada entrada soporta: visible, etiqueta i18n, ruta, icono, requerir-campaña, sólo-master.')
bullet(doc, 'Los cambios persisten en /users/me/preferences.')

h2(doc, '3.3 Tema y color de acento')
bullet(doc, 'Modo claro/oscuro/automatic (según hora del sistema).')
bullet(doc, 'Color de acento para acentos de UI.')

h2(doc, '3.4 URL base del backend')
bullet(doc, 'Texto para apuntar a un backend distinto (producción, staging, local).')
bullet(doc, 'Persistente en API_BASE_URL; aplica a todos los api clients.')

h2(doc, '3.5 Overlay de Skyline')
bullet(doc, 'Ajuste que define el tamaño máximo y posición por defecto del overlay skyline en la ventana de proyección.')

h2(doc, '3.6 Tamaño de ventanas secundarias')
bullet(doc, 'Define el tamaño por defecto al abrir windows de proyección (ancho x alto).')
bullet(doc, 'Persistente por usuario en /users/me/preferences.')

h2(doc, '3.7 Otros ajustes')
bullet(doc, 'Idioma de manuales (es | en) para cargar traducciones cuando la campaña permite varios manuales.')
bullet(doc, 'Selector de manuales activos por campaña: lista los manuales personalizados + dnd5e-2014.')
doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 4
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 4 · Pantallas secundarias (Mapas, Skyline, Proyección)')

h2(doc, '4.1 Visión general')
p(doc, 'Las páginas de proyección son rutas especiales del router que viven fuera del '
  'MainLayout protegido. Se sirven sin requerir JWT y exponen sólo lo permitido para jugadores.')

add_table(doc, ['Ruta', 'Página', 'Para quién', 'Características'], [
    ['/projection/maps', 'ProjectionMapPage', 'Jugadores', 'Mapa activo + grid + niebla de guerra + tokens aliados + luz + tiempo del día.'],
    ['/projection/skyline', 'ProjectionSkylinePage', 'Jugadores', 'Skyline activo + personaje en primer plano + overlay (título de pista sonora) + items skyline dinámicos.'],
], bold_first=False)

h2(doc, '4.2 Mapas')
p(doc, 'Los mapas son superficies rectangulares con imagen PNG/JPEG almacenada en el backend '
  '(backend sirve el binario en stream). Cada mapa puede tener hasta 4 imágenes variantes '
  '(una por momento del día: amanecer, día, atardecer, noche) y 4 skylines correspondientes.')

h3(doc, '4.2.1 Crear y gestionar mapas')
bullet(doc, 'Crear mapa: nombre (≤200 chars), dimensiones en mm o en píxeles, imagen subida.')
bullet(doc, 'Creación masiva: importMapToCampaign o /maps/bulk (varios a la vez).')
bullet(doc, 'Importar de otra campaña: ImportMapFromOtherCampaignDialog.')
bullet(doc, 'Toggle "preparado": marca el mapa como listo para usar desde el combobox de "mapa activo".')

h3(doc, '4.2.2 Grid Overlay')
bullet(doc, 'Activar/desactivar grid en /maps/:id y /projection/maps.')
bullet(doc, 'Tamaño de celda ajustable en píxeles.')
bullet(doc, 'Color y opacidad configurables.')
bullet(doc, 'Configurable por campaña desde /campaigns/:id/grid-overlay.')

h3(doc, '4.2.3 Niebla de guerra (Fog of War)')
bullet(doc, 'Modo cobertura: GLOBAL | por TOKEN | ninguno.')
bullet(doc, 'Modo pintura: REVEAL (mostrar) | HIDE (ocultar).')
bullet(doc, 'Modo pincel: square, circle.')
bullet(doc, 'Cuatro modos persistidos: completo, aliados solo, personalizada.')
bullet(doc, 'Niebla orgánica (curvas Bézier) complementaria a la niebla de cuadrícula.')

h3(doc, '4.2.4 Tokens')
bullet(doc, 'Cada token: id, posición (x, y), tamaño (S/M/L), color, etiqueta libre, sourceId (personaje o monstruo).')
bullet(doc, 'Modo edición: arrastrar, redimensionar, pegar etiqueta.')
bullet(doc, 'Modo lectura: popover al pulsar con resumen del personaje.')
bullet(doc, 'Tokens aliados se sincronizan con el Encounter activo (turno del personaje determina color del cuadro).')
bullet(doc, 'Sincronización ronda a ronda: useSkylineInitiativeSync repinta al cambiar turno.')

h3(doc, '4.2.5 Marcadores (Waypoints)')
bullet(doc, 'Crear marcador: nombre, posición, color, descripción.')
bullet(doc, 'Asociación opcional a monstruo/personaje/quest/worldpedia/diario/spell/shop.')
bullet(doc, 'Popup muestra resumen de la entidad asociada.')
bullet(doc, 'Para los jugadores, los marcadores se muestran pero son no interactivos.')

h3(doc, '4.2.6 Elementos del mapa')
add_table(doc, ['Tipo', 'Uso'], [
    ['MapWallElement', 'Bloquea línea de visión y movimiento.'],
    ['MapDoorElement', 'Permite paso condicional.'],
    ['MapWindowElement', 'Permite visión parcial.'],
    ['MapLightElement', 'Revela niebla en radio configurable.'],
    ['MapSoundSourceElement', 'Punto de sonido ambiental (canciones/SFX) con proximidad.'],
], bold_first=True)

h3(doc, '4.2.7 Skyline del mapa')
bullet(doc, 'Cuatro variantes por momento del día (amanecer, día, atardecer, noche).')
bullet(doc, 'Imagen PNG/JPEG almacenada en backend.')
bullet(doc, 'Se proyecta en /projection/maps pre-componiendo la imagen del mapa.')

h3(doc, '4.2.8 Importar / Exportar mapas')
bullet(doc, 'Importar mapa desde otra campaña: ImportMapFromOtherCampaignDialog.')
bullet(doc, 'Exportar datos del mapa como JSON: vía markers + elements + fog + tokens.')

h2(doc, '4.3 Skyline activo')
p(doc, 'Una "skyline" es la composición narrativa que se superpone al mapa: imagen del cielo, '
  'un personaje en primer plano, un título de pista sonora (overlay) y opcionalmente items '
  'flotantes (banderines, runas, etc.).')

h3(doc, '4.3.1 Default skyline por campaña')
bullet(doc, 'Una imagen por campaña, subida en /campaigns/:id/default-skyline.')
bullet(doc, 'Sirve como fondo cuando no hay skyline específica del mapa.')

h3(doc, '4.3.2 Skyline items')
bullet(doc, 'Items.addSkylineItem(name, side, positionX, positionY) para añadir overlays dinámicos.')
bullet(doc, 'Props slot (icono o texto) se muestra en la ventana de proyección.')
bullet(doc, 'Persistente entre sesiones hasta que se elimine.')

h3(doc, '4.3.3 Personaje en skyline')
bullet(doc, 'ActiveSkylineCharacter: el personaje activo en el combate se muestra en primer plano.')
bullet(doc, 'API: /campaigns/:id/active-skyline-character.')
bullet(doc, 'Visible en /projection/skyline, no en /projection/maps.')

h2(doc, '4.4 Ventanas Secundarias (Electron)')
bullet(doc, 'Proyección abre ventanas independientes sin chrome del navegador.')
bullet(doc, 'Tamaño configurable en /home → Ajustes.')
bullet(doc, 'Múltiples instancias permitidas.')
bullet(doc, 'Cierre automático al detener la sesión.')

doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 5
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 5 · Sidebar y tira del Master')
p(doc, 'El sidebar vive en MainLayout y combina navegación, reproductor de música/efectos, '
  'controles del tiempo del día, día actual del calendario y avisos.')

h2(doc, '5.1 Items del sidebar (configurables)')
add_table(doc, ['Key (SidebarItemDef)', 'Icono MUI', 'Etiqueta i18n', 'Ruta', 'Requiere campaña', 'Solo Master'], [
    ['campaigns', 'FolderSpecialIcon', 'Campañas', '/campaigns', 'No', 'No'],
    ['maps', 'MapIcon', 'Mapas', '/maps', 'Sí', 'No'],
    ['combate', 'SportsKabaddiIcon', 'Combate', '/combat', 'Sí', 'No'],
    ['soundtrack', 'MusicNoteIcon', 'Soundtrack', '/soundtrack', 'Sí', 'No'],
    ['soundEffects', 'BoltIcon', 'Efectos de sonido', '/soundtrack/effects', 'Sí', 'No'],
    ['characters', 'PeopleIcon', 'Personajes', '/characters', 'Sí', 'No'],
    ['bestiary', 'PetsIcon', 'Bestiario', '/campaign-bestiary', 'Sí', 'No'],
    ['spells', 'AutoFixHighIcon', 'Conjuros', '/campaign-spells', 'Sí', 'No'],
    ['manuals', 'MenuBookIcon', 'Manuales', '/manuals', 'No', 'No'],
    ['quests', 'AssignmentIcon', 'Misiones', '/quests', 'Sí', 'No'],
    ['shops', 'StorefrontIcon', 'Tiendas', '/shops', 'Sí', 'No'],
    ['worldpedia', 'AutoStoriesIcon', 'Worldpedia', '/worldpedia', 'Sí', 'No'],
    ['diary', 'EventNoteIcon', 'Diario', '/diary', 'Sí', 'No'],
    ['scenes', 'TheaterComedyIcon', 'Escenas', '/scenes', 'Sí', 'No'],
    ['cards', 'StyleIcon', 'Cartas', '/cards', 'No', 'No'],
], bold_first=False)

h2(doc, '5.2 Tira superior (AppBar)')
bullet(doc, 'Visible en pantallas estrechas (xs, sm).')
bullet(doc, 'Botón hamburguesa que abre el drawer temporal.')
bullet(doc, 'Nombre de la página actual (resolved via DEFAULT_SIDEBAR_ITEMS).')

h2(doc, '5.3 Time of Day Sidebar Controls')
bullet(doc, 'Selector rápido: Amanecer | Día | Atardecer | Noche.')
bullet(doc, 'Persistente en /campaigns/:id/time-of-day.')
bullet(doc, 'Aplica máscaras de luz y activa el audio apropiado.')

h2(doc, '5.4 Global Player (Soundtrack)')
h3(doc, '5.4.1 GlobalPlayerDrawerControls')
bullet(doc, 'Drawer accesible desde la esquina inferior izquierda.')
bullet(doc, 'Reproduce canciones del soundtrack global de la campaña.')
bullet(doc, 'Play/pause, siguiente, volumen, scrubbing.')
bullet(doc, 'Mantiene last-played, last queue, history.')

h3(doc, '5.4.2 Información de Now Playing')
bullet(doc, '"Now Playing" se publica en /soundtrack/projection/campaigns/:id/now-playing.')
bullet(doc, 'Visible en /projection/skyline vía SkylineOverlay.')
bullet(doc, 'Ocultable por SkylineOverlay toggle.')

h2(doc, '5.5 SFX Player')
bullet(doc, 'Reproductor dedicado para efectos de sonido cortos.')
bullet(doc, 'Hasta 5 efectos sonando en paralelo (cap configurable).')
bullet(doc, 'Auto-stop al cambiar mapa (useStopSfxOnMapChange).')
bullet(doc, 'Loop mode para efectos sostenidos (ambient, lluvia,etc).')

h2(doc, '5.6 ActiveScenesBar')
bullet(doc, 'Tira visible sólo para el Master en la parte inferior.')
bullet(doc, 'Botones one-click para reproducir las escenas recientes.')
bullet(doc, 'Detalle expandible al pulsar: opciones de la escena.')

h2(doc, '5.7 Shortcut Hotbar')
bullet(doc, 'Similar a ActiveScenesBar pero para atajos.')
bullet(doc, 'Pulsar un atajo lo ejecuta vía ShortcutRuntimeBridge (ver capítulo 7).')

h2(doc, '5.8 Diario en sidebar')
bullet(doc, 'Día actual del calendario de la campaña activa.')
bullet(doc, 'Navegación flecha izquierda/derecha para ir a días anterior/siguiente.')
bullet(doc, 'El master puede actualizar el día backend (updateCurrentDay).')

h2(doc, '5.9 Aviso: sesión activa')
bullet(doc, 'Si no hay sesión activa y la campaña la requiere, aparece un aviso interpolado.')
bullet(doc, 'Al pulsar redirige a /diary?tab=sessions&highlight=start con foco en "iniciar".')

h2(doc, '5.10 SkylinePreviewOverlay')
bullet(doc, 'Capa flotante que previsualiza una escena antes de dispararla.')
bullet(doc, 'Cierra con X o con tecla Escape.')

doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 6 — el más grueso
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 6 · Herramientas por módulo')

# 6.1 Campañas
h2(doc, '6.1 Campañas (/campaigns)')
h3(doc, '6.1.1 Crear campaña')
bullet(doc, 'Formulario (CampaignForm): nombre (≤200 chars), descripción (≤1000), idioma, color de acento, portada.')
bullet(doc, 'Portada subida a /campaigns/:id/image. Persistente en uploads.')
h3(doc, '6.1.2 Ajustes de campaña (CampaignSettingsModal)')
bullet(doc, 'Editar nombre, descripción, manuales activos, playlist de música.')
bullet(doc, 'Eliminar campaña (action irreversible).')
h3(doc, '6.1.3 Invitaciones')
bullet(doc, 'Lista de invitaciones pendientes con estado: pending | accepted | declined.')
bullet(doc, 'Re-invitar re-envía email.')
bullet(doc, 'Quitar jugador de la campaña: /campaigns/:id/player/:playerId.')
h3(doc, '6.1.4 Lista de campañas')
bullet(doc, 'Banner de campañas con cover, miembros, fecha de modificación.')
bullet(doc, 'Filtros: propias, compartidas, archivadas.')

# 6.2 Mapas
h2(doc, '6.2 Mapas (/maps)')
h3(doc, '6.2.1 Vista general: WorldMapView')
bullet(doc, 'Lista paginada con miniatura.')
bullet(doc, 'Toggle preparado (preparada ✓).')
bullet(doc, 'Acciones: editar, marcar activo, importar a otra campaña, eliminar.')

h3(doc, '6.2.2 Edición en vivo de un mapa')
bullet(doc, 'Image variant editor: subir imagen 1x por TOD.')
bullet(doc, 'Skyline editor: subir imagen 1x por TOD.')
bullet(doc, 'Grid overlay: celdas, color, opacidad.')
bullet(doc, 'Elementos: paredes, puertas, ventanas, luces, fuentes de sonido.')
bullet(doc, 'Marcadores (waypoints).')
bullet(doc, 'Niebla de guerra (cuadrícula y orgánica).')
bullet(doc, 'Tokens aliados sincronizados con el Encounter activo.')

h3(doc, '6.2.3 Importar mapas de otras campañas')
bullet(doc, 'ImportMapFromOtherCampaignDialog muestra todos los mapas que '
  'están en campañas donde el usuario fue miembro.')
bullet(doc, 'Seleccionar uno o varios y traerlos a la campaña activa.')

# 6.3 Soundtrack
h2(doc, '6.3 Soundtrack (/soundtrack)')
h3(doc, '6.3.1 Canciones')
bullet(doc, 'Subida de archivo MP3/OGG/WAV al backend (multipart).')
bullet(doc, 'Metadatos: nombre, autor, duración, tags (ambient, combate, descanso, misterio, etc.).')
bullet(doc, 'Streaming on-demand desde /soundtrack/songs/:id/stream.')
bullet(doc, 'Marcado de "played" para llevar al historial.')

h3(doc, '6.3.2 Playlists')
bullet(doc, 'Crear playlist con un nombre y selección de canciones.')
bullet(doc, 'Reordenar canciones vía drag (dnd-kit/sortable).')
bullet(doc, 'Reproducir playlist entera (cola secuencial).')

h3(doc, '6.3.3 Filtros')
bullet(doc, 'Filtros preestablecidos: ambient, combate, descanso, etc.')
bullet(doc, 'Cada filtro selecciona subconjunto de canciones Disponible en /soundtrack/filters.')

h3(doc, '6.3.4 Settings (SoundtrackSettingsPanel)')
bullet(doc, 'Toggle "Skyline muestra cancion actual": publicación en overlay.')
bullet(doc, 'Modo "encounter-driven": cambia automáticamente al iniciar encuentro.')
bullet(doc, 'Stop SFX al cambiar de mapa.')

h3(doc, '6.3.5 Historial (SongHistoryCard)')
bullet(doc, 'Lista cronológica de canciones reproducidas.')
bullet(doc, 'Botón "limpiar historial".')

# 6.4 Efectos de sonido
h2(doc, '6.4 Efectos de sonido (/soundtrack/effects)')
h3(doc, '6.4.1 Biblioteca de efectos')
bullet(doc, 'Subida de audio corto (≤30s recomendado).')
bullet(doc, 'Loop mode (none / loop / loop + crossfade).')
bullet(doc, 'Tags para filtros rápidos.')

h3(doc, '6.4.2 Presets')
bullet(doc, 'Preset agrupa varios efectos que se disparan a la vez.')
bullet(doc, 'Ej: preset "Puerta de taberna": pasos + murmullos + leña + algo de música.')
bullet(doc, 'Lista de presets por campaña + presets globales.')
bullet(doc, 'Editor visual: lista los items que se reproducirán en orden.')

h3(doc, '6.4.3 Source de sonido en mapa')
bullet(doc, 'MapSoundSourceElement con proximidad (useMapSoundPlayback).')
bullet(doc, 'Reproducción espacial cuando una fuente está dentro del rango de tokens aliados.')
bullet(doc, 'Reproducción automática al cambiar mapa (MapAudioOrchestrator).')

h3(doc, '6.4.4 Asociación a campañas')
bullet(doc, 'Effect.associate(unassociate) engancha un efecto a una campaña.')
bullet(doc, 'OnlyMaster solo efectos asociados a la campaña actual.')

# 6.5 Personajes
h2(doc, '6.5 Personajes (/characters)')
p(doc, 'El módulo de personajes cubre todo el ciclo de vida de un personaje de D&D 5e.')

h3(doc, '6.5.1 Lista de personajes (CharacterList)')
bullet(doc, 'Grid responsive con solo-thumbnail.')
bullet(doc, 'Filtros por nombre, clase, raza.')
bullet(doc, 'Acciones: abrir, editar, eliminar.')

h3(doc, '6.5.2 Hoja de personaje (CharacterDetailPage)')
bullet(doc, 'Abilities: STR, DEX, CON, INT, WIS, CHA con modificador.')
bullet(doc, 'Skills: lista completa con proficiencias y bonuses.')
bullet(doc, 'HP y AC con barras visuales.')
bullet(doc, 'Saving throws, senses (passive perception), speed.')
bullet(doc, 'Spell slots por nivel.')
bullet(doc, 'Lista de hechizos con preparable/conc.,cantrips, nivel y descripción.')
bullet(doc, 'Rasgos, dotes, inventario y notas del master.')

h3(doc, '6.5.3 Editor (CharacterEditorModal + CharacterAutoFillPanel)')
bullet(doc, 'Auto-relleno desde una clase y raza guardadas en la campaña.')
bullet(doc, 'Imagen: subida + crop circular o cuadrado.')
bullet(doc, 'Token image: crop cuadrado + subida.')
bullet(doc, 'AffinityLinks: relaciones con otros personajes (gráfico social).')

h3(doc, '6.5.4 AffinityChart')
bullet(doc, 'Visualización en grafo radial.')
bullet(doc, 'Tipos de relaciones: aliado, rival, familia, mentor, romántico.')
bullet(doc, 'Pulsar un nodo abre CharacterSheetModal inline.')

h3(doc, '6.5.5 Manuales de personajes (campaign-specific)')
bullet(doc, 'Personajes pueden referenciar: clase, raza, trasfondo, dote, rasgo, habilidad, conjuro, monstruo, manual.')
bullet(doc, 'Cada uno tiene su propia página CRUD (ver 6.6 a 6.14).')

# 6.6-6.13 Manuales de campaña
h2(doc, '6.6 Bestiario (/campaign-bestiary)')
bullet(doc, 'Lista paginada con thumbnails.')
bullet(doc, 'Crear monstruo: stat-block completo con HP, AC, ataques, traits, legendary actions.')
bullet(doc, 'Copiar desde manual: copyMonsterFromManual (dnd5e-2014 o personalizado).')
bullet(doc, 'Filtros por CR, tamaño, tipo.')
bullet(doc, 'Vista detallada: BestiaryDetailPage + MonsterStatBlock.')

h2(doc, '6.7 Conjuros (/campaign-spells)')
bullet(doc, 'Lista con nivel, escuela, tiempo de lanzamiento.')
bullet(doc, 'Editor: nivel, escuela (V,S,M components), casting time, range, duration, descripción, clases.')
bullet(doc, 'Importación/exportación masiva vía xlsx (/campaigns/:id/spells/import y /export).')
bullet(doc, 'Vista detallada: SpellStatBlock con descripción formateada.')

h2(doc, '6.8 Clases (/campaign-classes)')
bullet(doc, 'Editor de progresión por nivel: ClassLevelProgressionEditor.')
bullet(doc, 'Spell slots por nivel: ClassSpellcastingTable.')
bullet(doc, 'Features: lista de ClassFeature por nivel con descripción.')
bullet(doc, 'Copy desde manual dnd5e-2014.')
bullet(doc, 'Subir copy como plantilla para personalizar.')

h2(doc, '6.9 Razas (/campaign-races)')
bullet(doc, 'Subrazas y rasgos raciales.')
bullet(doc, 'Languages y proficiencias.')
bullet(doc, 'Speed, size, darkvision.')
bullet(doc, 'Vista detallada: stat-block compacto.')

h2(doc, '6.10 Trasfondos (/campaign-backgrounds)')
bullet(doc, 'Feature de trasfondo, ideal de características, posibles特技.')
bullet(doc, 'Vista detallada en /campaign-backgrounds y BestiaryDetailPage (extiende manual-detalles).')

h2(doc, '6.11 Dotes (/campaign-feats)')
bullet(doc, 'Editor: nombre, prerequisite, descripción, beneficios.')
bullet(doc, 'Stat-block compacto.')

h2(doc, '6.12 Rasgos (/campaign-traits)')
bullet(doc, 'Editor similar a feats pero sin prependiente.')

h2(doc, '6.13 Habilidades (/campaign-skills)')
bullet(doc, 'Editor de skill customization si difieren del manual base.')

# 6.14 Manuales
h2(doc, '6.14 Manuales (/manuals)')
h3(doc, '6.14.1 Manuales predefinidos')
bullet(doc, 'dnd5e-2014 (libros del SRD).')
bullet(doc, 'Cada manual expone TOC, búsqueda full-text, secciones.')
h3(doc, '6.14.2 Manuales personalizados (/manuals)')
bullet(doc, 'CRUD completo de manuales propios (custom-manuals).')
bullet(doc, 'Editor de entradas (ManualEditorPage):')
bullet(doc, 'Tipos de entrada: Background, Class, Monster, Race, Section, Simple, Spell.')
bullet(doc, 'Cover image: sube imagen de portada.')
bullet(doc, 'Editor markdown para Simple/Section.')
bullet(doc, 'Lista de entradas (EntryListPanel).')
h3(doc, '6.14.3 Importar / Exportar manuales')
bullet(doc, 'Export: descargar archivo JSON con todas las entradas.')
bullet(doc, 'Import: ImportManualDialog valida y crea el manual en la campaña.')

# 6.15 Misiones
h2(doc, '6.15 Misiones (/quests)')
bullet(doc, 'Lista: QuestList con QuestCard por misión.')
bullet(doc, 'Estados: activa | completada | fallida.')
bullet(doc, 'Editor: QuestFormDialog con título, descripción, estado, recompensas (texto libre).')
bullet(doc, 'Aperturas rápidas: vincular a worldpedia o a un mapa.')

# 6.16 Tiendas
h2(doc, '6.16 Tiendas (/shops)')
h3(doc, '6.16.1 Lista')
bullet(doc, 'Una Shop contiene Sections.')
bullet(doc, 'Cada Section contiene Columns y Entries.')
p(doc, 'Cada Entry tiene Cells de tipos:')
bullet(doc, 'TXT: texto editable inline (updateCellText).')
bullet(doc, 'MEDIA: imagen o data URI (uploadCellMedia) presentada como miniatura.')
bullet(doc, 'STREAM: streaming desde una fuente externa (getCellStreamUrl se usa para reproducir audio/video en la ventana del jugador).')
h3(doc, '6.16.2 Búsqueda')
bullet(doc, 'Búsqueda full-text sobre todas las entradas de la campaña.')
h3(doc, '6.16.3 Streaming para projection')
bullet(doc, 'Las Cells de tipo STREAM se exponen vía /shops/cells/:cellId/stream y pueden empujarse a /projection/skyline como items dinámicos.')

# 6.17 Diario
h2(doc, '6.17 Diario (/diary)')
h3(doc, '6.17.1 Calendario')
bullet(doc, 'Configurable por campaña (DiaryCalendarSettings):')
bullet(doc, 'Meses custom (nombre y días).')
bullet(doc, 'Día actual (currentDay) modificable vía flechas y backend.')
bullet(doc, 'Fecha visualizada en sidebar.')
h3(doc, '6.17.2 Sesiones')
bullet(doc, 'Lista de sesiones: lista DiariSessions con startSesion,endSesion,deleteSesion.')
bullet(doc, 'Una sola sesión activa a la vez.')
bullet(doc, 'visit-day marca cada día tocado en el diario.')
h3(doc, '6.17.3 Entradas diarias')
bullet(doc, 'Editor rich text (Quill) por día.')
bullet(doc, 'Upsert (crear o reemplazar) por día.')
bullet(doc, 'Auto-link a personas/quest/worldpedia del mismo día.')

# 6.18 Worldpedia
h2(doc, '6.18 Worldpedia (/worldpedia)')
h3(doc, '6.18.1 Árbol jerárquico')
bullet(doc, 'Folders > Notes. Drag&drop para reordenar.')
bullet(doc, 'Carpetas pueden anidarse (WorldpediaMoveDialog).')
h3(doc, '6.18.2 Notas (WorldpediaNoteEditor)')
bullet(doc, 'Editor rich text (Quill) con soporte Markdown.')
bullet(doc, 'Imágenes con auth (/api/auth-image).')
bullet(doc, 'Auto-links a monstruos/hechizos/personajes/quest/shop/mapas (loadAutoLinkRules).')
h3(doc, '6.18.3 Backlinks')
bullet(doc, 'WorldpediaBacklinks muestra qué otras notas apuntan a esta.')
h3(doc, '6.18.4 Búsqueda')
bullet(doc, 'WorldpediaSearchBar full-text sobre todas las notas (titles + contents).')
h3(doc, '6.18.5 Export / Import')
bullet(doc, 'Export de una nota, una carpeta o todo el árbol (JSON).')
bullet(doc, 'Import sobreescribe o agrega respetando la jerarquía.')

# 6.19 Cartas
h2(doc, '6.19 Cartas (/cards)')
p(doc, 'Generador de cartas imprimibles (dados, hechizos, rasgos, etc.) basado en plantillas.')

h3(doc, '6.19.1 Plantillas de carta')
bullet(doc, 'Editor visual (CardTemplateEditorDialog) con:')
bullet(doc, 'Selector de presets de tamaño (POKER, MINI, BRIDGE, TAROT, LETTER, CUSTOM).')
bullet(doc, 'Orientación (portrait/landscape).')
bullet(doc, 'Slots arrastrables: TEXT_SINGLE, TEXT_MULTI, IMAGE, KEY_VALUE_LIST, DIVIDER, FRAME, BADGE.')
bullet(doc, 'Color, grosor, bordes personalizados por slot.')
bullet(doc, 'Bindeo a campos de entidad (name, level, school, etc).')
bullet(doc, 'Botones de rotación, flip, alineación, match-size, distribute.')

h3(doc, '6.19.2 Generador de cartas (CharacterCardGeneratorDialog)')
bullet(doc, 'Selecciona personaje + plantilla.')
bullet(doc, 'Auto-empaqueta hechizos/rasgos/dotes seleccionados en cartas.')
bullet(doc, 'Vista previa de todas las cartas en grid.')
bullet(doc, 'Export: PDF (jsPDF + html2canvas) o imprimir desde navegador (printCardsViaBrowser).')
bullet(doc, 'Page format preset: A4, A3, A5, etc. (paperFormats.ts).')
bullet(doc, 'Multi-carta-por-página automático.')

h3(doc, '6.19.3 Slots DIVIDER (decorativo)')
bullet(doc, 'Tipos de efecto: plano, cadena, cuerda, fuego, hilos.')
bullet(doc, 'Grosor configurable en mm.')
bullet(doc, 'Curvatura perpendicular (curveMm).')
bullet(doc, 'End taper asimétrico.')
bullet(doc, 'Color configurable desde picker dedicado.')

h3(doc, '6.19.4 Herramientas del editor')
bullet(doc, 'Zoom + pan con flecha-arriba/abajo/izquierda/derecha.')
bullet(doc, 'Pan por click & drag sobre el canvas.')
bullet(doc, 'Snap to grid (alineación automática con otro slot).')

# 6.20 Escenas (scenes)
h2(doc, '6.20 Escenas (/scenes)')
p(doc, 'Composiciones multimedia reproducibles disparadas por el Master.')

h3(doc, '6.20.1 Tipos de elementos')
bullet(doc, 'Audio: archivo MP3/OGG/WAV con play/pause/stop.')
bullet(doc, 'Video: archivo MP4 con play/pause/stop y seek.')
bullet(doc, 'Narrador TTS: síntesis de voz (roboti formant engine).')
bullet(doc, 'Imagen: PNG/JPEG con fade in/out.')

h3(doc, '6.20.2 Síntesis de voz (Roboti)')
bullet(doc, 'Formant engine custom (frontend/src/components/scenes/utils/narrator/robotiFormantEngine.ts).')
bullet(doc, 'Continuous stream arquitecture plan (ROBOTI_CONTINUOUS_STREAM_BLUEPRINT).')
bullet(doc, 'Spanish benchmark phrases integradas.')
bullet(doc, 'Configurables: voiceConfig.roboti con F1-F4,BW1-BW4,F0,aspiration,etc.')
bullet(doc, 'Pitch, velocidad, duración estimada.')
bullet(doc, 'Cache por (text, voice config, synth version).')

h3(doc, '6.20.3 Parking en mapa')
bullet(doc, 'Una escena estacionada se reproduce automáticamente al activar el mapa.')

h3(doc, '6.20.4 ActiveScenesBar (ver 5.6)')
bullet(doc, 'Acceso rápido a escenas recientes.')

# 6.21 Atajos
h2(doc, '6.21 Atajos (/shortcuts)')
p(doc, 'Sistema de comandos programables con macros. (Ver capítulo 7 para detalle.)')

# 6.22 Combate
h2(doc, '6.22 Combate (/combat)')
h3(doc, '6.22.1 Lista de encuentros')
bullet(doc, 'EncounterList con EncounterCard.')
bullet(doc, 'Estado: activo, inactivo.')
h3(doc, '6.22.2 Editor de encuentro (EncounterFormDialog)')
bullet(doc, 'Lista de aliados (personajes) y enemigos (monstruos o personajes).')
bullet(doc, 'Recursos: tope inicial, dificultad calculada (computeEncounterMetrics).')
h3(doc, '6.22.3 CombatView')
bullet(doc, 'InitiativePanel: orden de iniciativa con round/turn.')
bullet(doc, 'ParticipantsPanel: lista de jugadores y enemigos.')
bullet(doc, 'CombatNotesBox: notas rápidas durante el combate.')
bullet(doc, 'CombatHeader: header visual con HP, AC.')
h3(doc, '6.22.4 Proyección de combate')
bullet(doc, 'CombatView escucha useSkylineInitiativeSync: al cambiar turno, '
  'los tokens aliados se sincronizan con /projection/maps.')
bullet(doc, 'Encuentroactivo se publica en /campaigns/:id/active-encounter.')

# 6.23 Sonido en mapa
h2(doc, '6.23 MapAudioOrchestrator')
bullet(doc, 'Componente headless montado en MainLayout.')
bullet(doc, 'Reproduce automáticamente la canción apropiada según mapa activo + tiempo del día.')
bullet(doc, 'Aplica filtros de acuerdo al TOD (noche = más graves, día = completo).')
bullet(doc, 'useStopSfxOnMapChange garantiza que los SFX del mapa anterior se apagan.')

# 6.24 Tokens aliados en mapa
h2(doc, '6.24 Tokens aliados sincronizados')
bullet(doc, 'useMapTokens lee tokens del backend y reparte por Encounter activo.')
bullet(doc, 'useSkylineInitiativeSync empuja los tokens aliados al /projection/maps.')
bullet(doc, 'TokenQuickInfoPopover muestra resumen al pulsar.')

# 6.25 Proyecciones de skyline
h2(doc, '6.25 Window/SkylineOverlay')
bullet(doc, 'Overlay opcional que muestra el título de la canción actual en la proyección.')
bullet(doc, 'Ajustable per-campaign desde /home o via /campaigns/:id/skyline-overlay.')

# 6.26 Update checker (en home)
h2(doc, '6.26 Actualizaciones automáticas (electron-updater)')
bullet(doc, 'Comprueba versión en intervalo predefinido.')
bullet(doc, 'Modal con changelog y botón "Instalar y reiniciar".')

# 6.27 Misc
h2(doc, '6.27 Otros componentes y helpers')
bullet(doc, 'invitationsList: banner con invitaciones pendientes en MainLayout.')
bullet(doc, 'DebugUserInfo: usado en desarrollo.')
bullet(doc, 'TitleBar: chrome personalizado de Electron.')

doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 7
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 7 · Atajos (Shortcuts)')
p(doc, 'Shortcuts es un sistema de macros disparadas por eventos. Cada Shortcut está compuesto '
  'por una lista de Actions que se ejecutan en cadena.')

h2(doc, '7.1 Estructura de un atajo')
bullet(doc, 'Nombre y descripción.')
bullet(doc, 'Icono visual.')
bullet(doc, 'Acciones encadenadas, cada una con sus parámetros.')

h2(doc, '7.2 Tipos de acciones disponibles')
add_table(doc, ['Acción', 'Efecto'], [
    ['playSong', 'Reproduce una canción del soundtrack.'],
    ['playEffect', 'Reproduce un SFX.'],
    ['stopAllAudio', 'Detiene soundtrack y SFX.'],
    ['setTimeOfDay', 'Cambia el tiempo del día.'],
    ['setActiveMap', 'Cambia el mapa activo.'],
    ['setActiveEncounter', 'Cambia el encuentro activo.'],
    ['playScene', 'Reproduce una escena.'],
    ['sendEmote', 'Envía un emote a la skyline.'],
    ['showOverlay', 'Activa overlay de skyline.'],
    ['runTTS', 'Ejecuta narrador TTS.'],
], bold_first=True)

h2(doc, '7.3 Disparadores')
bullet(doc, 'Botón manual (ShortcutHotbar / SidebarShortcutsPanel).')
bullet(doc, 'Hotkey global (Configuración → Atajos globales → combinación de teclas).')
bullet(doc, 'URL shortcut (deep-linking).')

h2(doc, '7.4 Editor visual')
bullet(doc, 'Drag&drop para reordenar acciones.')
bullet(doc, 'Cada acción muestra sus parámetros en un form.')
bullet(doc, 'Validación cliente: si la acción requiere una canción, se muestra la combobox.')

h2(doc, '7.5 Persistencia')
bullet(doc, 'Backend: /shortcuts/:id persistido en Postgres (o SQLite).')
bullet(doc, 'Frontend: ShortcutsContext expone array y operaciones CRUD.')
bullet(doc, 'Runtime: ShortcutRuntimeBridge intercepta hotkeys.')

doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 8
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 8 · Iconos y ventanas flotantes Skyline')

h2(doc, '8.1 SkylinePreviewOverlay')
bullet(doc, 'Overlay que aparece al previsualizar una escena antes de dispararla.')
bullet(doc, 'Tamaño reducible / maximizable.')
bullet(doc, 'Cierra con X o Escape.')

h2(doc, '8.2 ActiveScenesBar')
bullet(doc, 'Botones one-click a escenas recientes.')
bullet(doc, 'Cada botón reproduce esa escena específica.')

h2(doc, '8.3 ShortcutHotbar')
bullet(doc, 'Hotbar ubicada en el pie del MainLayout.')
bullet(doc, 'Hasta 8 botones visibles.')
bullet(doc, 'Reordenable.')

h2(doc, '8.4 SkylineFloatingItems')
bullet(doc, 'Items temporales que se renderizan sobre el mapa (banderines, runas).')
bullet(doc, 'Persiste vía /campaigns/skyline-items.')

h2(doc, '8.5 SkylineOverlay')
bullet(doc, 'Texto de canción actual en la parte superior de la proyección.')
bullet(doc, 'Color y tipografía configurables.')

h2(doc, '8.6 ActiveSkylineCharacterBadge')
bullet(doc, 'Muestra retrato de personaje activo en combate sobre la skyline.')

doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 9
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 9 · Funcionalidades adicionales (Importación, Exportación, etc.)')

h2(doc, '9.1 Importación')
add_table(doc, ['Tipo', 'Origen', 'Ruta'], [
    ['Hechizos (xlsx)', 'Archivo .xlsx pre-formateado', '/campaigns/:id/spells/import'],
    ['Monstruos (manual)', 'Manual dnd5e-2014 o custom', 'copyMonsterFromManual'],
    ['Clases/Razas/Trasfondos/Dotes/Rasgos/Habilidades', 'Manual', 'copyClassFromManual / copyRaceFromManual / etc.'],
    ['Mapas desde otras campañas', 'Campaña donde el usuario fue miembro', 'importMapToCampaign'],
    ['Manuales', 'Archivo .json exportado', '/custom-manuals/import'],
    ['Worldpedia', 'Archivo .json exportado', '/worldpedia/campaigns/:id/import'],
], bold_first=True)

h2(doc, '9.2 Exportación')
add_table(doc, ['Tipo', 'Formato'], [
    ['Hechizos', 'Excel .xlsx vía /campaigns/:id/spells/export'],
    ['Worldpedia (nota)', 'JSON vía /worldpedia/campaigns/:id/export/notes/:id'],
    ['Worldpedia (carpeta)', 'JSON vía /worldpedia/campaigns/:id/export/folders/:id'],
    ['Worldpedia (todo)', 'JSON vía /worldpedia/campaigns/:id/export'],
    ['Manuales', 'JSON vía /custom-manuals/:id/export'],
    ['Mapas (imagen)', 'PNG/JPEG descargable desde /maps/:id/image'],
    ['Mapas (skydata)', 'JSON de markers + elements + fog + tokens'],
], bold_first=True)

h2(doc, '9.3 Tamaños de ventanas secundarias')
bullet(doc, 'Ajustable desde /home → Tamaño de ventanas secundarias.')
bullet(doc, 'Persistente por usuario.')
bullet(doc, 'useSecondaryWindowSizes aplica al abrir BrowserWindow.')

h2(doc, '9.4 Update checker')
bullet(doc, 'electron-updater busca nueva versión cada 24 h.')
bullet(doc, 'Release notes visibles en diálogo.')
bullet(doc, 'Instalación con "Instalar y reiniciar".')

h2(doc, '9.5 QR de campaña')
bullet(doc, 'qrcode.react usado para generar QR en /campaigns/:id/share.')
bullet(doc, 'Permite invitar escaneando desde móvil.')

h2(doc, '9.6 Crop de imagen')
bullet(doc, 'react-easy-crop como librería de UI.')
bullet(doc, 'Aplica a: foto de personaje, token image, scene image.')
bullet(doc, 'sharp en backend para redimensionar.')

h2(doc, '9.7 i18n')
bullet(doc, 'i18next con namespaces (cards, manual, soundtrack, etc).')
bullet(doc, 'Idiomas disponibles: es, en.')
bullet(doc, 'Locale auto-seleccionado del navegador.')
bullet(doc, 'Cambio manual desde /home.')

doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 10
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 10 · Acceso web, Electron y multi-dispositivo')

h2(doc, '10.1 Modos de despliegue')
add_table(doc, ['Modo', 'Tecnología', 'Notas'], [
    ['Web', 'SPA en navegador', 'Acceso vía http(s), requiere backend accesible.'],
    ['Electron desktop', 'electron 38 + electron-builder', 'Instalador .exe/.dmg/.AppImage/.deb/.rpm.'],
    ['Update server', 'electron-updater + release/', 'Canal de auto-actualización opcional.'],
], bold_first=True)

h2(doc, '10.2 Ventanas de proyección (Electron)')
bullet(doc, 'Apertura desde /maps o /scenes con el botón "Proyectar".')
bullet(doc, 'BrowserWindow sin chrome y dimensiones configurables.')
bullet(doc, 'Cierre automático al apagar la sesión.')

h2(doc, '10.3 Sync entre ventanas')
add_table(doc, ['Mecanismo', 'Uso'], [
    ['BroadcastChannel', 'Sync rápida entre tabs del mismo origen.'],
    ['LocalStorage event', 'Backup de BroadcastChannel.'],
    ['REST polling', 'Sync master ↔ projection windows (~1-2s polling).'],
    ['SocketIO', 'Sockets bidireccionales; en uso parcial (battle-state).'],
], bold_first=True)

h2(doc, '10.4 Backend compartido')
bullet(doc, 'Todas las ventanas consultan el mismo backend.')
bullet(doc, 'JWT en localStorage.')
bullet(doc, 'Endpoints públicos (/campaigns/projection/:id/*) sin autenticación.')

h2(doc, '10.5 Modo offline')
bullet(doc, 'No soportado actualmente. La app requiere conexión al backend.')
bullet(doc, 'Cache de Audio y datos pesados: see src/api/cache.')

doc.add_page_break()

# ---------------------------------------------------------------------------
# Capítulo 11
# ---------------------------------------------------------------------------
h1(doc, 'Capítulo 11 · Roles de usuario: Master vs Jugador')

h2(doc, '11.1 Master')
p(doc, 'El Master es el dueño de la campaña o un co-Master invitado. Tiene acceso completo '
  'a la edición de contenido y al control de la proyección.')

h3(doc, '11.1.1 Privilegios')
bullet(doc, 'Crear, editar, eliminar campañas.')
bullet(doc, 'Crear, editar, eliminar monstruos, hechizos, clases, etc.')
bullet(doc, 'Editar mapas (niebla, grid, marcadores, elementos).')
bullet(doc, 'Iniciar encuentros, gestionar iniciativa.')
bullet(doc, 'Disparar atajos, reproducir soundtrack y efectos.')
bullet(doc, 'Ajustar Time of Day, Active Map, Active Encounter.')
bullet(doc, 'Cambiar manual(es) activos de la campaña.')
bullet(doc, 'Importar / exportar contenido.')

h2(doc, '11.2 Jugador (Player)')
p(doc, 'El Jugador es un miembro de la campaña con status accepted. Su acceso es limitado '
  'y de sólo lectura en casi todas las superficies.')

h3(doc, '11.2.1 Privilegios')
bullet(doc, 'Leer contenido de campaña (manuales, bestiario, hechizos, etc).')
bullet(doc, 'Editar únicamente su propio personaje (/characters/:id con permisos).')
bullet(doc, 'Abrir ventanas de proyección: /projection/maps y /projection/skyline.')
bullet(doc, 'Ver la skyline y now-playing overlay.')

h2(doc, '11.3 Invitado (pending)')
bullet(doc, 'Tiene invitación emitida pero no aceptada.')
bullet(doc, 'Puede aceptar y rechazar.')
bullet(doc, 'Tras aceptar, su rol (master/player) determina los privilegios.')

h2(doc, '11.4 Co-Master')
bullet(doc, "Master con rol=\"co\" (extensión futura).")
bullet(doc, 'Por ahora, todos los miembros master tienen los mismos permisos.')

doc.add_page_break()

# ---------------------------------------------------------------------------
# Apéndice A — Rutas
# ---------------------------------------------------------------------------
h1(doc, 'Apéndice A · Mapa de rutas (frontend router)')
p(doc, 'Lista completa de rutas expuestas por el RouterProvider (createHashRouter).')

add_table(doc, ['Ruta', 'Página', 'Protección'], [
    ['/projection/maps', 'ProjectionMapPage', 'Público (projection-only).'],
    ['/projection/skyline', 'ProjectionSkylinePage', 'Público (projection-only).'],
    ['/', 'HomePage', 'ProtectedLayout + MainLayout (root redirect).'],
    ['/change-password', 'ChangePasswordPage', 'ProtectedLayout.'],
    ['/delete-account', 'DeleteAccountPage', 'ProtectedLayout.'],
    ['/campaigns', 'CampaignPage', 'ProtectedLayout.'],
    ['/soundtrack', 'SoundtrackPage', 'ProtectedLayout + camp.'],
    ['/soundtrack/effects', 'SoundEffectsPage', 'ProtectedLayout + camp.'],
    ['/maps', 'MapsPage', 'ProtectedLayout + camp.'],
    ['/shortcuts', 'ShortcutsPage', 'ProtectedLayout.'],
    ['/combat', 'CombatPage', 'ProtectedLayout + camp.'],
    ['/characters', 'CharactersPage', 'ProtectedLayout + camp.'],
    ['/characters/:id', 'CharacterDetailPage', 'ProtectedLayout + camp.'],
    ['/diary', 'DiaryPage', 'ProtectedLayout + camp. + DiarySidebarContext.'],
    ['/quests', 'QuestsPage', 'ProtectedLayout + camp.'],
    ['/shops', 'ShopsPage', 'ProtectedLayout + camp.'],
    ['/scenes', 'ScenesPage', 'ProtectedLayout + camp.'],
    ['/worldpedia', 'WorldpediaPage', 'ProtectedLayout + camp.'],
    ['/campaign-bestiary', 'CampaignBestiaryPage', 'ProtectedLayout + camp.'],
    ['/campaign-spells', 'CampaignSpellsPage', 'ProtectedLayout + camp.'],
    ['/cards', 'CardsPage', 'ProtectedLayout.'],
    ['/manuals/:manualId/edit', 'ManualEditorPage', 'ProtectedLayout.'],
    ['/manuals', 'ManualsHomePage', 'Público (es).'],
    ['/manuals/:manualId', 'ManualViewerPage', 'Público.'],
    ['/manuals/:manualId/section/:nodeId', 'ManualViewerPage', 'Público.'],
    ['/manuals/:manualId/classes/:id', 'ClassPage', 'Público.'],
    ['/spells', 'SpellsPage', 'Público.'],
    ['/login', 'LoginPage', 'Público.'],
    ['/register', 'RegisterPage', 'Público.'],
    ['/forgot-password', 'ForgotPasswordPage', 'Público.'],
    ['/reset-password', 'ResetPasswordPage', 'Público.'],
], bold_first=True)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Apéndice B — API surface
# ---------------------------------------------------------------------------
h1(doc, 'Apéndice B · Superficie de la API (endpoints backend)')

h2(doc, 'B.1 Módulos autenticados (/auth, /users)')
add_table(doc, ['Método', 'Ruta', 'Handler'], [
    ['POST', '/auth/register', 'register'],
    ['POST', '/auth/login', 'login'],
    ['POST', '/auth/forgot-password', 'forgotPassword'],
    ['POST', '/auth/reset-password', 'resetPassword'],
    ['PUT', '/auth/change-password', 'changePassword'],
    ['GET', '/users/me', 'me'],
    ['PATCH', '/users/me/preferences', 'preferences'],
    ['DELETE', '/users/me', 'deleteAccount'],
    ['GET', '/users/:id', 'findOne'],
    ['GET', '/network-info', 'LAN IP discovery'],
], bold_first=True)

h2(doc, 'B.2 Campañas (/campaigns)')
add_table(doc, ['Método', 'Ruta', 'Notas'], [
    ['POST/GET', '/campaigns', 'CRUD básico.'],
    ['GET/PATCH/DELETE', '/campaigns/:id', 'Por id, master-only.'],
    ['POST', '/campaigns/:id/invite', 'Invitación por email.'],
    ['GET', '/campaigns/invitations/pending', 'Lista invitaciones del usuario.'],
    ['POST', '/campaigns/invitation/respond', 'Aceptar/rechazar.'],
    ['DELETE', '/campaigns/:id/player/:playerId', 'Quitar jugador.'],
    ['GET/PATCH', '/campaigns/:id/active-encounter', 'Set/get active encounter.'],
    ['GET/PATCH', '/campaigns/:id/active-map', 'Active map.'],
    ['GET/PATCH', '/campaigns/:id/active-skyline-character', 'Active character on skyline.'],
    ['GET/PATCH', '/campaigns/:id/battle-state', 'Battle state (HP, status, turn).'],
    ['GET/PATCH', '/campaigns/:id/default-skyline', 'Default skyline image.'],
    ['GET/PATCH', '/campaigns/:id/fog-of-war', 'Default fog state.'],
    ['GET/PATCH', '/campaigns/:id/grid-overlay', 'Default grid config.'],
    ['GET/PATCH', '/campaigns/:id/manuals', 'Manuals activos por campaña.'],
    ['GET/POST/DELETE', '/campaigns/:id/skyline-items', 'Skyline floating items.'],
    ['GET/PATCH', '/campaigns/:id/skyline-overlay', 'Skyline overlay settings.'],
    ['GET/PATCH', '/campaigns/:id/soundtrack-settings', 'Soundtrack settings.'],
    ['GET/PATCH', '/campaigns/:id/time-of-day', 'TOD aktivo.'],
], bold_first=True)

h2(doc, 'B.3 Proyección pública (/campaigns/projection/:id)')
add_table(doc, ['Método', 'Ruta', 'Notas'], [
    ['GET', '/campaigns/projection/:id/battle-state', 'Battle estado (público).'],
    ['GET', '/campaigns/projection/:id/default-skyline', 'Default skyline image (público).'],
    ['GET', '/campaigns/projection/:id/default-skyline/exists', 'Bool.'],
    ['GET', '/campaigns/projection/:id/participant-monster-map', 'Map participante→monstruo.'],
    ['GET', '/campaigns/projection/:id/skyline-overlay', 'Overlay settings (público).'],
], bold_first=True)

h2(doc, 'B.4 Manuales (/manuals + /custom-manuals)')
bullet(doc, 'GET /manuals — lista de manuales.')
bullet(doc, 'GET /manuals/:id/toc — tabla de contenidos.')
bullet(doc, 'GET /manuals/:id/sections/:nodeId — sección por id.')
bullet(doc, 'GET /manuals/:id/search — búsqueda full-text.')
bullet(doc, 'GET /manuals/:id/spells | /spells/:id | /spells/meta/all.')
bullet(doc, 'GET /manuals/:id/monsters | /monsters/:slug.')
bullet(doc, 'GET /manuals/:id/classes | /classes/:id.')
bullet(doc, 'GET /manuals/:id/races | /races/:id.')
bullet(doc, 'GET /manuals/:id/traits | /traits/:id.')
bullet(doc, 'GET /manuals/:id/feats | /feats/:id.')
bullet(doc, 'GET /manuals/:id/skills | /skills/:id.')
bullet(doc, 'GET /manuals/:id/backgrounds | /backgrounds/:id.')
bullet(doc, 'CRUD /custom-manuals.* y /custom-manuals/:id/entries.*.')
bullet(doc, 'GET /custom-manuals/:id/export | POST /custom-manuals/import.')

h2(doc, 'B.5 Contenido de campaña')
add_table(doc, ['Entidad', 'Endpoints'], [
    ['Monstruos', 'CRUD /campaigns/:id/bestiary + copy/manual + export/import'],
    ['Conjuros', 'CRUD /campaigns/:id/spells + copy/manual + xlsx export/import'],
    ['Clases', 'CRUD /campaigns/:id/classes + copy/manual'],
    ['Razas', 'CRUD /campaigns/:id/races + copy/manual'],
    ['Trasfondos', 'CRUD /campaigns/:id/backgrounds + copy/manual'],
    ['Dotes', 'CRUD /campaigns/:id/feats + copy/manual'],
    ['Rasgos', 'CRUD /campaigns/:id/traits + copy/manual'],
    ['Habilidades', 'CRUD /campaigns/:id/skills + copy/manual'],
    ['Encounters', 'CRUD /campaigns/:id/encounters'],
], bold_first=True)

h2(doc, 'B.6 Mapas (/maps)')
add_table(doc, ['Método', 'Ruta', 'Notas'], [
    ['POST/GET', '/maps', 'CRUD básico, bulk upload.'],
    ['GET/POST/PATCH/DELETE', '/maps/:id', 'Por id.'],
    ['GET/POST', '/maps/:id/image', 'Upload/servicio imagen.'],
    ['GET/POST', '/maps/:id/skyline', 'Upload/servicio skyline.'],
    ['GET/PATCH', '/maps/:id/elements', 'Map elements (walls, doors...).'],
    ['GET/PATCH', '/maps/:id/fog', 'Fog cuandrícula.'],
    ['GET/PATCH', '/maps/:id/organic-fog', 'Fog orgánico.'],
    ['GET/PATCH', '/maps/:id/tokens', 'Tokens aliados.'],
    ['GET/POST/PATCH/DELETE', '/maps/:id/markers', 'Marcadores (waypoints).'],
    ['PATCH', '/maps/:id/prepared', 'Toggle preparado.'],
    ['POST', '/maps/:id/import', 'Import legacy map data.'],
    ['GET', '/maps/usage', 'Estadísticas de uso.'],
    ['GET', '/maps/other-campaigns', 'Listar mapas de otras campañas.'],
], bold_first=True)

h2(doc, 'B.7 Soundtrack (/soundtrack + /soundtrack/effects + /soundtrack/presets)')
add_table(doc, ['Método', 'Ruta', 'Notas'], [
    ['GET/POST', '/soundtrack/songs', 'CRUD canciones.'],
    ['GET', '/soundtrack/campaigns/:id/songs', 'Catálogo por campaña.'],
    ['GET', '/soundtrack/songs/:id/stream', 'Stream audio.'],
    ['POST', '/soundtrack/songs/:id/played', 'Marca como reproducida (historial).'],
    ['POST/DELETE', '/soundtrack/songs/:id/associate(/:campaignId)', 'Asociación a campaña.'],
    ['GET/POST/PATCH/DELETE', '/soundtrack/campaigns/:id/playlists', 'CRUD playlist.'],
    ['GET/PATCH', '/soundtrack/presets/campaigns/:id', 'Presets por campaña.'],
    ['GET/POST/PATCH/DELETE', '/soundtrack/effects', 'CRUD efectos cortos.'],
    ['POST/DELETE', '/soundtrack/effects/:id/associate(/:campaignId)', 'Asociación.'],
    ['GET', '/soundtrack/usage', 'Histórico agregago.'],
    ['GET', '/soundtrack/projection/campaigns/:id/now-playing', 'Público.'],
], bold_first=True)

h2(doc, 'B.8 Diario (/diary)')
add_table(doc, ['Método', 'Ruta', 'Notas'], [
    ['GET/PATCH', '/diary/campaigns/:id/calendar', 'Config calendario.'],
    ['GET/PATCH', '/diary/campaigns/:id/calendar/current-day', 'Día actual.'],
    ['GET/POST', '/diary/campaigns/:id/entries', 'CRUD entradas.'],
    ['GET', '/diary/campaigns/:id/entries/:year/:month/:day', 'Por día.'],
    ['GET', '/diary/campaigns/:id/entries/by-id/:id', 'Por id.'],
    ['POST', '/diary/campaigns/:id/entries/upsert', 'Upsert por día.'],
    ['GET/POST/PATCH/DELETE', '/diary/campaigns/:id/sessions', 'Sesiones.'],
    ['GET', '/diary/campaigns/:id/sessions/active', 'Sesión activa.'],
    ['POST', '/diary/campaigns/:id/sessions/:id/start', 'Iniciar sesión.'],
    ['POST', '/diary/campaigns/:id/sessions/:id/end', 'Terminar sesión.'],
    ['POST', '/diary/campaigns/:id/sessions/:id/visit-day', 'Registrar visita.'],
], bold_first=True)

h2(doc, 'B.9 Worldpedia (/worldpedia)')
add_table(doc, ['Método', 'Ruta', 'Notas'], [
    ['GET', '/worldpedia/campaigns/:id/tree', 'Árbol completo.'],
    ['GET', '/worldpedia/campaigns/:id/search', 'Búsqueda full-text.'],
    ['POST/PATCH/DELETE', '/worldpedia/campaigns/:id/folders', 'CRUD carpetas.'],
    ['POST/PATCH/DELETE', '/worldpedia/campaigns/:id/notes', 'CRUD notas.'],
    ['PATCH', '/worldpedia/campaigns/:id/notes/:id/move', 'Mover nota.'],
    ['GET', '/worldpedia/campaigns/:id/notes/:id/links', 'Backlinks.'],
    ['PATCH', '/worldpedia/campaigns/:id/reorder', 'Drag&drop persistence.'],
    ['GET', '/worldpedia/campaigns/:id/export', 'Export completo (JSON).'],
    ['GET', '/worldpedia/campaigns/:id/export/folders/:id', 'Export carpeta.'],
    ['GET', '/worldpedia/campaigns/:id/export/notes/:id', 'Export nota.'],
    ['POST', '/worldpedia/campaigns/:id/import', 'Import.'],
], bold_first=True)

h2(doc, 'B.10 Otros')
add_table(doc, ['Módulo', 'Endpoints'], [
    ['Quests', 'CRUD /quests'],
    ['Shops', 'CRUD /shops + nested /shops/:id/sections, /sections/:id/columns, /sections/:id/entries, /entries/:id/cells.'],
    ['Characters', 'CRUD /characters + /affinity-links.'],
    ['Spells (global)', 'GET /spells, /spells/meta/all.'],
    ['Maps projection', 'No aplica — projection lee battle-state por /campaigns/projection/:id/battle-state.'],
], bold_first=True)

doc.add_page_break()

# ---------------------------------------------------------------------------
# Apéndice C — Tech stack
# ---------------------------------------------------------------------------
h1(doc, 'Apéndice C · Stack tecnológico')
p(doc, 'Resumen de paquetes por scope. Lista completa y versiones en docs/TECH_STACK.md.')

add_table(doc, ['Scope', 'Categoría', 'Paquetes'], [
    ['Frontend', 'UI', '@mui/material 7, @mui/icons-material 7, @mui/x-data-grid 8, @emotion/react 11'],
    ['Frontend', 'Core', 'react 18, react-dom, react-router-dom 6, vite 5'],
    ['Frontend', 'DnD', '@dnd-kit/core 6, /sortable 10, /utilities 3'],
    ['Frontend', 'i18n', 'i18next 25, react-i18next 16'],
    ['Frontend', 'Red', 'axios 1, socket.io-client 4'],
    ['Frontend', 'Rich content', 'quill 1, react-quill 2, react-markdown 9, remark-gfm 4'],
    ['Frontend', 'Media', 'react-easy-crop 5, qrcode.react 4'],
    ['Frontend', 'XLSX', 'xlsx 0.18'],
    ['Frontend', 'Tipos', 'typescript 5'],
    ['Frontend', 'Util', 'sass-embedded, @types/uuid, uuid 13'],
    ['Backend', 'Core', '@nestjs/* 10, typescript 5'],
    ['Backend', 'ORM', 'typeorm 0.3, better-sqlite3 12, sqlite3, reflect-metadata'],
    ['Backend', 'Auth', 'passport, passport-jwt, @nestjs/jwt, @nestjs/passport, bcrypt'],
    ['Backend', 'Red', '@nestjs/websockets, @nestjs/platform-socket.io 10, socket.io 4'],
    ['Backend', 'Docs/validación', '@nestjs/swagger 7, class-validator 0.14, class-transformer 0.5'],
    ['Backend', 'Imagen', 'sharp 0.33'],
    ['Backend', 'Email', 'resend 6'],
    ['Backend', 'XLSX', 'xlsx 0.18'],
    ['Backend', 'Sanitización', 'sanitize-html 2'],
    ['Electron', 'Shell', 'electron 38, electron-builder 26, electron-updater 6'],
    ['Tooling', 'Lint', 'eslint, prettier, @typescript-eslint'],
    ['Tooling', 'Test backend', 'jest 29, supertest, ts-jest'],
], bold_first=True)

h2(doc, 'C.1 Tecnologías implícitas')
bullet(doc, 'Web Audio API (audio espacial/proximidad en mapa).')
bullet(doc, 'BroadcastChannel + localStorage (sync entre vistas/pestañas).')
bullet(doc, 'Electron IPC (sync ventanas secundarias).')
bullet(doc, 'Socket.IO (battle state, fog sync).')
bullet(doc, 'Quill (rich text).')
bullet(doc, 'React Router (hash router en Electron para file://).')

doc.add_page_break()

# ---------------------------------------------------------------------------
# Apéndice D — Glosario y atajos globales
# ---------------------------------------------------------------------------
h1(doc, 'Apéndice D · Glosario y atajos globales del teclado')

h2(doc, 'D.1 Glosario de términos')
add_table(doc, ['Término', 'Definición'], [
    ['Campaña', 'Contenedor raíz de una mesa D&D: jugadores, manuales, configuración.'],
    ['Manuales', 'Bases de reglas (D&D 5e o personalizadas) que proveen contenido base.'],
    ['Encuentro', 'Lista de aliados/enemigos en un combate con iniciativa y orden de turnos.'],
    ['Skyline', 'Composición narrativa que se proyecta a los jugadores.'],
    ['SFX', 'Efecto de sonido corto (puerta,espada,grito,etc.).'],
    ['Soundtrack', 'Música ambiente y canciones de fondo.'],
    ['Fog of War', 'Niebla de guerra: oculta porciones del mapa.'],
    ['Grid Overlay', 'Rejilla que se dibuja sobre el mapa.'],
    ['Marcador (waypoint)', 'Punto interesable en el mapa con link a entidad.'],
    ['TOD', 'Time of day: amanecer, día, atardecer, noche.'],
    ['Worldpedia', 'Wiki jerárquica Markdown de la campaña.'],
    ['Carta', 'Tarjeta imprimible generada con plantilla + entidad.'],
    ['Escena', 'Composición multimedia reproducible.'],
    ['Atajo', 'Macro encadenada de acciones disparada por hotkey o botón.'],
    ['Affinity chart', 'Grafo de relaciones entre personajes.'],
    ['CRUD', 'Create / Read / Update / Delete.'],
    ['Active Encounter/Map/Skyline Character', 'Estado actual de la sesión.'],
    ['Battle state', 'Estado de HP, status, turno, ronda del encuentro activo.'],
    ['Hovered/Selected token', 'Token resaltado actualmente.'],
    ['Proyección', 'Vista del jugador con mapa o skyline.'],
    ['Recurring actor', 'Personaje que se reutiliza entre escenas.'],
    ['Editor narración', 'Síntesis de voz vía Roboti formant engine.'],
], bold_first=True)

h2(doc, 'D.2 Atajos globales del teclado (frontend)')
add_table(doc, ['Atajo', 'Función'], [
    ['Ctrl/Cmd + D', 'Duplicar slot seleccionado en editor de plantillas (cards).'],
    ['Esc', 'Cerrar overlay de previsualización Skyline.'],
    ['Ctrl/Cmd + Wheel', 'Zoom en previsualización del editor de plantillas.'],
    ['Arriba/Abajo/Izq/Der en cards', 'Pan del canvas sin scroll horizontal.'],
    ['] y [', 'Subir / bajar capa del slot seleccionado.'],
    ['Click&drag canvas', 'Pan (en cards).'],
], bold_first=True)

h2(doc, 'D.3 Notas de implementación')
bullet(doc, 'Los atajos arriba son los globales; cada Shortcut (cap. 7) puede tener su propio hotkey configurable.')
bullet(doc, 'Las teclas son sensibles a Repeat; useKeyEvent.preventDefault() en cada atajo.')
bullet(doc, 'Las hotkeys están deshabilitadas cuando el foco está en campos de texto.')


# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

doc.save(OUT_PATH)
print(f'Saved {OUT_PATH}')
